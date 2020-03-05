#!/user/  .in Unix only

#  ------------------------------------
#  Porject: LEARN
#   Author: Rainer Warth
#  Version: 28-02-2020 Previous: 04.04.2016 First: 04.04.2016
#    Goals: Test programm to use the docx package
#      Ref: e-book "Do borring things withpyhton" chapter 13
#      Ref: 
#      Ref: 
#    Satus: <runs> - <bug (false output , script does not run)> - <broken (link, module, file is missing)> 
#    Satus: broken
#       >N: --

'''
Description: Test programm to use the docx package
Source:     e-book "Do borring things withpyhton" chapter 13
PyVersion:	3.4
Date:		
>DO:
'''

import docx
doc = docx.Document('demo.docx')

# doc = docx.Document("C:\\Users\\Public\\Documents\\Pub_CodeDev\\TestData\\160407_MSWordFiles\\n-160415_Verwaltungsleiter_MPI-Freiburg\\n-160415_All-Verwaltungsleiter_MPI-Freiburg.docx")

print(len(doc.paragraphs))
