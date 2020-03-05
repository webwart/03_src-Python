#!/user/  .in Unix only

#  ------------------------------------
#  Porject: LEARN
#   Author: Rainer Warth
#  Version: 28-02-2020 Previous: 05.04.2016 First: 05.04.2016
#    Goals: Find MSWord file in directory, Open file, Find table, Print table
#      Ref: https://pathlib.readthedocs.org/en/pep428/
#      Ref: https://docs.python.org/3/library/pathlib.html?highlight=pathlib#module-pathlib
#    Satus: <runs> - <bug (false output , script does not run)> - <broken (link, module, file is missing)> 
#    Satus: broken
#       >N: --

import docx
from pathlib import Path, PureWindowsPath
allTables = docx.Document()
#1 path to MS word files and make list all files to use

PathDir = "C:/Users/Public/Documents/Pub_CodeDev/TestData/160407_MSWordFiles"
p = Path(PathDir)
# print([x for x in p.iterdir() if x.is_dir()])
for directory in p.iterdir():
    if directory.is_dir():
        print("\n::: New directory :::")
        print(directory)
        wordFiles = (list(directory.glob('**/*_All-*.docx')))
        for wordFile in wordFiles:
            print("::::: WordFile in WordFiles :::::")
            '''
            f = open(str(PureWindowsPath(wordFile)))    #-- DOES NOT WORK, but would be nice
            doc = docx.Document(f)
            # print(len(doc.paragraphs))
            f.close()
'''
#2 open file
            pathStringWordFile = (PureWindowsPath(wordFile))
            print(str(pathStringWordFile))
            doc = docx.Document(str(pathStringWordFile))
            print(len(doc.tables))
            print(doc.tables[0])
            # allTables.save('160407_AllTables.docx') 
            # >N: Find the table paragraph
            
        

#3 read the first paragraphs
# doc.paragraphs[0].text
# doc.add_paragraph('Hello world!')
# doc.save('helloworld.docx')

# moresophisticated
#4 select paragraph with table
#5 store tabel in an table object
#6 print table object
