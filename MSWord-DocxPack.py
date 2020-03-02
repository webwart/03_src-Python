#!/user/  .in Unix only

#  ------------------------------------
#  Porject: LEARN
#   Author: Rainer Warth
#  Version: 28-02-2020
#    Goals: Deal with MS-Word documents 
#      Ref: --
#    Satus: <runs> - <bug (false output , script does not run)> - <broken (link, module, file is missing)> 
#    Satus: broken
#       >N: --
#  ------------------------------------


from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

'''
document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)
'''
# document.add_picture('monty-truth.png', width=Inches(1.25))

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
'''for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item.qty)
    row_cells[1].text = str(item.id)
    row_cells[2].text = item.desc
'''
document.add_page_break()

document.save('TestDocxOutput.docx')
